# resume_parser.py
"""
简历解析模块

本模块负责从不同格式的文件中提取文本内容：
1. PDF文档解析（包括OCR功能）
2. Word文档(DOCX)解析
3. HTML内容转纯文本
4. 文本清理和格式化
"""

import tempfile
import os
import logging
import re
import hashlib
import time
from bs4 import BeautifulSoup
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import docx
from io import BytesIO

def compact_resume_text(text: str) -> str:
    # Simple trimming implementation
    return " ".join(text.split())

def md5_hash(text: str) -> str:
    return hashlib.md5(text.encode("utf-8")).hexdigest()

def html_to_text(html_str: str) -> str:
    """优化的HTML转文本函数，更好地处理简历格式"""
    if not html_str:
        return ""
    soup = BeautifulSoup(html_str, "lxml")
    # 删除脚本和样式
    for element in soup(["script", "style", "head"]):
        element.decompose()
    # 处理换行
    for br in soup.find_all("br"):
        br.replace_with("\n")
    # 处理段落
    for p in soup.find_all("p"):
        p.append(soup.new_string("\n"))
    text = soup.get_text(separator=" ")
    # 清理多余空白
    text = re.sub(r'\s*\n\s*', '\n', text)
    text = re.sub(r' +', ' ', text)
    return text.strip()

def parse_pdf(file_data: bytes) -> str:
    """
    解析PDF文件并提取文本内容
    
    实现了多种文本提取策略:
    1. 直接提取PDF文本
    2. 当直接提取不可用时，使用OCR识别页面
    3. 提取和识别PDF中的嵌入图片
    
    Args:
        file_data: PDF文件的二进制数据
        
    Returns:
        str: 提取的文本内容
    """
    if not file_data:
        logging.error("PDF数据为空")
        return ""
        
    temp_name = None
    doc = None
    text_parts = []  # Initialize text_parts list
    
    try:
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            temp_name = temp_file.name
            temp_file.write(file_data)
            temp_file.flush()
            
        doc = fitz.open(temp_name)
        for page_num in range(len(doc)):
            try:
                page = doc[page_num]
                # Try regular text extraction first
                text = page.get_text()
                if text.strip():
                    text_parts.append(text)
                    continue
                
                # If no text found, try OCR
                pix = page.get_pixmap()
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                img_text = pytesseract.image_to_string(img, lang='chi_sim+eng')
                if img_text.strip():
                    text_parts.append(img_text)
                
                # Try to extract and OCR images in the page
                for img_info in page.get_images():
                    try:
                        xref = img_info[0]
                        base_image = doc.extract_image(xref)
                        if base_image:
                            image = Image.open(BytesIO(base_image["image"]))
                            img_text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                            if img_text.strip():
                                text_parts.append(img_text)
                    except Exception as e:
                        logging.warning(f"处理PDF页面{page_num}中的图片失败: {e}")
                        continue
                        
            except Exception as e:
                logging.error(f"处理PDF页面{page_num}失败: {e}")
                continue
                
    except Exception as e:
        logging.error(f"PDF解析失败: {e}")
        return ""
        
    finally:
        if doc:
            try:
                doc.close()
            except Exception:
                pass
        if temp_name and os.path.exists(temp_name):
            try:
                os.unlink(temp_name)
            except Exception:
                pass

    # Process and clean the extracted text
    if text_parts:
        # Join all parts and clean up
        text = "\n".join(text_parts)
        lines = [line.strip() for line in text.split("\n")]
        clean_lines = []
        prev_line = None
        for line in lines:
            if line and line != prev_line:  # Remove empty and duplicate lines
                clean_lines.append(line)
                prev_line = line
        return "\n".join(clean_lines)
    
    return ""  # Return empty string if no text was extracted

def extract_text_from_pdf_ocr(file_data: bytes) -> str:
    """使用OCR提取PDF中的文字"""
    try:
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            temp_name = temp_file.name
            temp_file.write(file_data)
            temp_file.flush()
        
        doc = fitz.open(temp_name)
        text_parts = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = pytesseract.image_to_string(img, lang='chi_sim+eng')
            text_parts.append(text)
        
        doc.close()
        os.unlink(temp_name)
        return "\n".join(text_parts).strip()
    except Exception as e:
        logging.error(f"PDF OCR解析失败: {e}")
        return ""

def parse_docx(file_data: bytes) -> str:
    """解析DOCX文件并返回文本内容"""
    try:
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_name = temp_file.name
            temp_file.write(file_data)
            temp_file.flush()
        
        doc = docx.Document(temp_name)
        text_parts = []
        
        # 提取段落文本
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 处理段落样式
                if para.style.name.startswith('Heading'):
                    text = f"\n{text}\n"
                text_parts.append(text)
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    # 获取单元格中的所有段落
                    cell_text = '\n'.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    text_parts.append(' | '.join(row_texts))
        
        # 处理文档中的图片
        try:
            from docx.shape import InlineShape
            for shape in doc.inline_shapes:
                if shape.type == InlineShape.PICTURE:
                    # 保存图片到临时文件
                    img_byte_array = shape._inline.graphic.graphicData.pic.blipFill.blip.embed._blob
                    image = Image.open(BytesIO(img_byte_array))
                    # OCR处理图片
                    img_text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                    if img_text.strip():
                        text_parts.append(img_text)
        except Exception as e:
            logging.error(f"处理Word文档中的图片失败: {e}")
        
        os.unlink(temp_name)
        
        # 合并所有文本并清理
        text = "\n".join(text_parts)
        # 删除重复行和多余空白
        lines = text.split("\n")
        clean_lines = []
        prev_line = None
        for line in lines:
            line = line.strip()
            if line and line != prev_line:
                clean_lines.append(line)
                prev_line = line
        return "\n".join(clean_lines)
    except Exception as e:
        logging.error(f"DOCX解析失败: {e}")
        return ""
